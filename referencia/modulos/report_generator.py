import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from fpdf import FPDF
from datetime import datetime
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from email.mime.image import MIMEImage  # <-- Importante: nueva importación
import os


def exportar_a_excel(
    datos_crudos, df_procesado, nombre_archivo="reporte_criptomonedas.xlsx"
):
    """
    Exporta los datos crudos y procesados a un archivo de Excel con dos hojas.

    Args:
        datos_crudos (list): La lista de diccionarios con los datos originales de la API.
        df_procesado (pd.DataFrame): El DataFrame con los datos limpios y calculados.
        nombre_archivo (str): El nombre del archivo Excel a generar.
    """
    if not datos_crudos or df_procesado is None:
        print("No hay datos para exportar a Excel.")
        return

    try:
        df_crudo = pd.DataFrame(datos_crudos)
        with pd.ExcelWriter(nombre_archivo, engine="openpyxl") as writer:
            # Hoja 1: Datos Crudos
            df_crudo.to_excel(writer, sheet_name="Datos Crudos", index=False)

            # Hoja 2: Cálculos y Datos Limpios
            df_procesado.to_excel(writer, sheet_name="Cálculos", index=False)

        print(f"\nDatos exportados exitosamente a '{nombre_archivo}'")
    except Exception as e:
        print(f"Ocurrió un error al exportar a Excel: {e}")


def generar_reporte_word(df_procesado, nombre_archivo="reporte_tendencias.docx"):
    """Genera un resumen detallado con análisis de tendencias en un documento de Word."""
    if df_procesado is None or df_procesado.empty:
        print("No hay datos para generar el reporte de Word.")
        return
    try:
        activo_mayor_ganancia = df_procesado.loc[
            df_procesado["Variacion % (24h)"].idxmax()
        ]
        activo_mayor_perdida = df_procesado.loc[
            df_procesado["Variacion % (24h)"].idxmin()
        ]
        variacion_promedio = df_procesado["Variacion % (24h)"].mean()
        activos_al_alza = (df_procesado["Variacion % (24h)"] > 0).sum()
        activos_a_la_baja = (df_procesado["Variacion % (24h)"] < 0).sum()
        total_activos = len(df_procesado)

        doc = Document()
        doc.add_heading("Análisis de Tendencias de Criptomonedas", level=0)
        fecha_actual = datetime.now().strftime("%d de %B de %Y, %H:%M:%S")
        doc.add_paragraph(f"Reporte generado el: {fecha_actual}")

        doc.add_heading("Resumen General del Mercado (Últimas 24h)", level=1)
        sentimiento = "positivo" if variacion_promedio > 0 else "negativo"
        p_resumen = doc.add_paragraph(
            f"El análisis de los {total_activos} principales activos muestra un sentimiento de mercado generalmente {sentimiento} en las últimas 24 horas, con una variación porcentual promedio de "
        )
        p_resumen.add_run(f"{variacion_promedio:.2f}%.").bold = True
        doc.add_paragraph(
            f"Activos con variación positiva: {activos_al_alza}", style="List Bullet"
        )
        doc.add_paragraph(
            f"Activos con variación negativa: {activos_a_la_baja}", style="List Bullet"
        )

        doc.add_heading("Activos Destacados", level=1)
        doc.add_heading("Mayor Ganancia", level=2)
        p_ganador = doc.add_paragraph("El activo con el mejor rendimiento fue ")
        p_ganador.add_run(f"{activo_mayor_ganancia['Activo']}").bold = True
        p_ganador.add_run(f", con un impresionante aumento del ")
        p_ganador.add_run(f"{activo_mayor_ganancia['Variacion % (24h)']:.2f}%").bold = (
            True
        )
        p_ganador.add_run(
            f", alcanzando un precio de ${activo_mayor_ganancia['Precio Actual (USD)']:,.2f}."
        )

        doc.add_heading("Mayor Pérdida", level=2)
        p_perdedor = doc.add_paragraph(
            "Por otro lado, el activo con el rendimiento más bajo fue "
        )
        p_perdedor.add_run(f"{activo_mayor_perdida['Activo']}").bold = True
        p_perdedor.add_run(f", con una caída del ")
        p_perdedor.add_run(f"{activo_mayor_perdida['Variacion % (24h)']:.2f}%").bold = (
            True
        )
        p_perdedor.add_run(
            f", situando su precio en ${activo_mayor_perdida['Precio Actual (USD)']:,.2f}."
        )

        doc.add_heading("Tabla de Datos Detallada", level=1)
        table = doc.add_table(rows=1, cols=df_procesado.shape[1], style="Table Grid")
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(df_procesado.columns):
            hdr_cells[i].text = col_name
        for index, row in df_procesado.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row["Activo"])
            row_cells[1].text = f"${row['Precio Actual (USD)']:,.2f}"
            row_cells[2].text = f"{row['Variacion % (24h)']:.2f}%"

        doc.save(nombre_archivo)
        print(f"\nReporte de Word generado exitosamente en '{nombre_archivo}'")
    except Exception as e:
        print(f"Ocurrió un error al generar el documento de Word: {e}")


def guardar_grafico_como_imagen(df, nombre_imagen="variacion_criptos.png"):
    """
    Genera el gráfico de barras y lo guarda como un archivo de imagen.

    Args:
        df (pd.DataFrame): DataFrame con los datos procesados.
        nombre_imagen (str): Nombre del archivo de imagen a guardar.

    Returns:
        str: El nombre del archivo de la imagen guardada.
    """
    if df is None or df.empty:
        print("No hay datos para generar el gráfico.")
        return None

    plt.figure(figsize=(12, 7))
    colors = ["green" if x > 0 else "red" for x in df["Variacion % (24h)"]]
    bars = plt.bar(df["Activo"], df["Variacion % (24h)"], color=colors)
    plt.xlabel("Criptomoneda")
    plt.ylabel("Variación Porcentual (%) en 24h")
    plt.title("Variación Porcentual de las Principales Criptomonedas en 24h")
    plt.xticks(rotation=45, ha="right")

    for bar in bars:
        yval = bar.get_height()
        va_align = "bottom" if yval >= 0 else "top"
        plt.text(
            bar.get_x() + bar.get_width() / 2.0,
            yval,
            f"{yval:.2f}%",
            va=va_align,
            ha="center",
        )

    plt.grid(axis="y", linestyle="--", alpha=0.7)
    plt.tight_layout()

    # En lugar de plt.show(), guardamos la figura
    plt.savefig(nombre_imagen)
    plt.close()  # Cerramos la figura para liberar memoria
    print(f"\nGráfico guardado como '{nombre_imagen}'")
    return nombre_imagen


class PDF(FPDF):
    def header(self):
        # Este método se deja vacío para tener control total sobre la portada y las páginas de contenido.
        pass

    def footer(self):
        # Posición a 1.5 cm del final
        self.set_y(-15)
        # Fuente Arial itálica 8
        self.set_font("Helvetica", "I", 8)
        # Número de página
        self.cell(0, 10, "Página " + str(self.page_no()), 0, 0, "C")


# VERSIÓN CORREGIDA DE LA FUNCIÓN
def generar_reporte_pdf(
    df_procesado, ruta_grafico, nombre_archivo="reporte_ejecutivo.pdf"
):
    """
    Genera un reporte ejecutivo en PDF con portada, resumen y gráfico.
    YA NO ELIMINA EL GRÁFICO.
    """
    if df_procesado is None or not ruta_grafico:
        print("Faltan datos o el gráfico para generar el PDF.")
        return

    try:
        # --- 1. Análisis para el Mini-Resumen ---
        activo_mayor_ganancia = df_procesado.loc[
            df_procesado["Variacion % (24h)"].idxmax()
        ]
        activo_mayor_perdida = df_procesado.loc[
            df_procesado["Variacion % (24h)"].idxmin()
        ]
        variacion_promedio = df_procesado["Variacion % (24h)"].mean()

        # --- 2. Creación del PDF ---
        pdf = PDF()  # Asumiendo que tienes la clase PDF definida

        # --- Portada ---
        pdf.add_page()
        # ... (resto del código de la portada) ...
        pdf.set_font("Helvetica", "B", 24)
        pdf.cell(0, 80, "", ln=True)
        pdf.cell(0, 10, "Reporte Ejecutivo", ln=True, align="C")
        # ... etc.

        # --- Página de Contenido ---
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 15, "Resumen para la Gerencia", ln=True)

        pdf.set_font("Helvetica", "", 11)
        resumen_texto = (
            f"El análisis de las últimas 24 horas muestra una tendencia general de mercado con una variación promedio de {variacion_promedio:.2f}%. "
            f"El activo con mejor desempeño fue {activo_mayor_ganancia['Activo']}, registrando un alza de {activo_mayor_ganancia['Variacion % (24h)']:.2f}%. "
            f"En contraste, {activo_mayor_perdida['Activo']} fue el activo con la mayor caída, disminuyendo un {activo_mayor_perdida['Variacion % (24h)']:.2f}%. "
            "Se recomienda monitorear la volatilidad y los activos destacados."
        )
        pdf.multi_cell(0, 6, resumen_texto)

        pdf.ln(10)
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 15, "Rendimiento de Activos (Últimas 24h)", ln=True)
        pdf.image(ruta_grafico, x=15, w=180)

        # --- 3. Guardar el PDF ---
        pdf.output(nombre_archivo)
        print(f"\nReporte PDF generado exitosamente en '{nombre_archivo}'")

    except Exception as e:
        print(f"Ocurrió un error al generar el PDF: {e}")


def enviar_correo_con_adjuntos(
    lista_destinatarios, archivos_adjuntos, datos_plantilla, ruta_grafico
):
    """
    Envía un correo HTML con un gráfico y tabla embebidos, además de los archivos adjuntos.
    """
    smtp_server = "smtp.gmail.com"
    smtp_port = 587
    email_user = "mail.com"  # Cambia por tu correo
    email_password = "pass"  # Tu contraseña de aplicación

    # --- Creación del Mensaje ---
    # El contenedor principal es 'mixed' para combinar contenido y adjuntos.
    msg = MIMEMultipart("mixed")
    msg["From"] = email_user
    msg["To"] = ", ".join(lista_destinatarios)
    msg["Subject"] = (
        f"Reporte de Mercado de Criptomonedas - {datos_plantilla['{FECHA}']}"
    )

    # Contenedor 'related' para el cuerpo HTML y la imagen embebida.
    msg_related = MIMEMultipart("related")

    # --- Cuerpo del Correo (HTML) ---
    try:
        with open("plantilla.html", "r", encoding="utf-8") as f:
            html = f.read()

        for key, value in datos_plantilla.items():
            html = html.replace(key, str(value))

        # Adjuntar el HTML al contenedor 'related'
        msg_alternative = MIMEMultipart("alternative")
        msg_alternative.attach(MIMEText(html, "html"))
        msg_related.attach(msg_alternative)

    except FileNotFoundError:
        print("Error: No se encontró 'plantilla.html'.")
        return  # Salir si no hay plantilla

    # --- Embeber el Gráfico ---
    try:
        with open(ruta_grafico, "rb") as f_img:
            img = MIMEImage(f_img.read())
            # El 'Content-ID' debe coincidir con el 'cid:' en el HTML
            img.add_header("Content-ID", "<grafico_variacion>")
            msg_related.attach(img)
    except FileNotFoundError:
        print(
            f"Advertencia: No se encontró el gráfico '{ruta_grafico}'. El correo se enviará sin la imagen embebida."
        )

    # Adjuntar el contenedor 'related' (HTML + imagen) al mensaje principal
    msg.attach(msg_related)

    # --- Adjuntar Archivos (Excel, Word, PDF) ---
    for archivo_path in archivos_adjuntos:
        try:
            with open(archivo_path, "rb") as attachment:
                part = MIMEApplication(
                    attachment.read(), Name=os.path.basename(archivo_path)
                )
                part["Content-Disposition"] = (
                    f'attachment; filename="{os.path.basename(archivo_path)}"'
                )
                msg.attach(part)
        except FileNotFoundError:
            print(f"Advertencia: No se encontró el adjunto '{archivo_path}'.")

    # --- Envío del Correo ---
    try:
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(email_user, email_password)
        server.sendmail(email_user, lista_destinatarios, msg.as_string())
        print("\nCorreo con reportes enriquecidos enviado exitosamente.")
    except Exception as e:
        print(f"Error al enviar el correo: {e}")
    finally:
        if "server" in locals():
            server.quit()
