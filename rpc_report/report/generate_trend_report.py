import logging
import pandas as pd
from docx import Document
from datetime import datetime

log = logging.getLogger(__name__)


def generate_trend_report(df: pd.DataFrame, filename="trend_report.docx"):
    """
    Genera un resumen detallado con análisis de tendencias en un documento de Word.

    Args:
        df (pd.DataFrame): El DataFrame con los datos limpios y calculados.
        filename (str): El nombre del archivo Word a generar.
    """

    if df is None or df.empty:
        log.info("No hay datos para generar el reporte de tendencias en Word.")
        return

    try:
        # Variaciones de las principales criptomonedas
        activo_mayor_ganancia = df.loc[df["price_change_percentage_24h"].idxmax()]
        activo_mayor_perdida = df.loc[df["price_change_percentage_24h"].idxmin()]
        variacion_promedio = df["price_change_percentage_24h"].mean()

        # Contar activos con variación positiva y negativa
        activos_al_alza = (df["price_change_percentage_24h"] > 0).sum()
        activos_a_la_baja = (df["price_change_percentage_24h"] < 0).sum()
        total_activos = len(df)

        # Generar el documento de Word
        # Introducción
        doc = Document()
        doc.add_heading("Análisis de Tendencias de Criptomonedas", level=0)
        fecha_actual = datetime.now().strftime("%d de %B de %Y, %H:%M:%S")
        doc.add_paragraph(f"Reporte generado el: {fecha_actual}")

        # Resumen General
        doc.add_heading("Resumen General del Mercado (Últimas 24h)", level=1)
        sentimiento = "positivo" if variacion_promedio > 0 else "negativo"
        p_resumen = doc.add_paragraph(
            f"El análisis de los {total_activos} principales activos muestra un sentimiento de mercado generalmente {sentimiento} en las últimas 24 horas, con una variación porcentual promedio de "
        )
        p_resumen.add_run(f"{variacion_promedio:.2f}%.").bold = True
        doc.add_paragraph(
            f"Activos con variación positiva: {activos_al_alza}", style="List Bullet"
        )
        doc.add_paragraph(
            f"Activos con variación negativa: {activos_a_la_baja}", style="List Bullet"
        )

        # Activos Destacados mayor ganancia
        doc.add_heading("Activos Destacados", level=1)
        doc.add_heading("Mayor Ganancia", level=2)
        p_ganador = doc.add_paragraph("El activo con el mejor rendimiento fue ")
        p_ganador.add_run(f"{activo_mayor_ganancia['name']}").bold = True
        p_ganador.add_run(f", con un impresionante aumento del ")
        p_ganador.add_run(
            f"{activo_mayor_ganancia['price_change_percentage_24h']:.2f}%"
        ).bold = True
        p_ganador.add_run(
            f", alcanzando un precio de ${activo_mayor_ganancia['current_price']:,.2f}."
        )

        # Activos Destacados más devaluado
        doc.add_heading("Mayor Pérdida", level=2)
        p_perdedor = doc.add_paragraph(
            "Por otro lado, el activo con el rendimiento más bajo fue "
        )
        p_perdedor.add_run(f"{activo_mayor_perdida['name']}").bold = True
        p_perdedor.add_run(f", con una caída del ")
        p_perdedor.add_run(
            f"{activo_mayor_perdida['price_change_percentage_24h']:.2f}%"
        ).bold = True
        p_perdedor.add_run(
            f", situando su precio en ${activo_mayor_perdida['current_price']:,.2f}."
        )

        # Tabla de Datos
        doc.add_heading("Tabla de Datos Detallada", level=1)
        table = doc.add_table(rows=1, cols=df.shape[1], style="Table Grid")
        hdr_cells = table.rows[0].cells

        for i, col_name in enumerate(df.columns):
            hdr_cells[i].text = col_name

        for index, row in df.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row["name"])
            row_cells[1].text = f"${row['current_price']:,.2f}"
            row_cells[2].text = f"{row['price_change_percentage_24h']:.2f}%"

        doc.save(filename)
        log.info(f"Reporte generado exitosamente: '{filename}'")
    except Exception as e:
        log.error(f"Error al generar el reporte: {e}")
